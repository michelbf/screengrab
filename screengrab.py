from selenium import webdriver

from docx import Document
from docx.shared import Inches

import time
import re
import os
import argparse

#====================

msg_mode = """
*** Screen Grabber ***
Usage:  usage: screengrab.py [-h] [--headless] [-c <chromedriver_path>] mode file

[--headless]   -   Optional. Runs Chrome Driver in background.
file           -   Specify file where addresses are. 

Your source file addresses should be in the <ip>:<port> format, or in the http(s)://something.com format.

IMPORTANT: If you don't provide the -c argument. ChromeDriver needs to be downloaded 
           to your machine and in your PATH.
*********
"""

parser=argparse.ArgumentParser(usage=msg_mode)
parser.add_argument("--headless", help="Runs Chrome Driver in background.", action="store_true")
parser.add_argument("-c", "--chromedriver", help="Path to ChromeDriver", default='chromedriver')
parser.add_argument("file", help="File containing the IPs:Port or Domain Names")
args=parser.parse_args()

#====================

#Initializing webdriver:
DRIVER=args.chromedriver
chrome_options = webdriver.ChromeOptions()
chrome_options.add_argument("--no-sandbox")
chrome_options.add_argument("--disable-dev-shm-usage")
if (args.headless):
	chrome_options.add_argument("--headless")

driver = webdriver.Chrome(DRIVER,chrome_options=chrome_options)
driver.implicitly_wait(20)
driver.set_page_load_timeout(20)

#====================

def create_dir():
	dirName = "./" + str(int(time.time()))
	if (not os.path.isdir(dirName)):
		os.mkdir(dirName)
	return dirName


def read_file(fname):
    with open(fname) as f:
        lines = set(f.read().splitlines())
    return lines



def checkFormat(s):
	IP_REGEX = r"(\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)):(\d{1,5}\b)"
	URL_REGEX=r"^(http:\/\/www\.|https:\/\/www\.|http:\/\/|https:\/\/)?[a-z0-9]+([\-\.]{1}[a-z0-9]+)*\.[a-z]{2,5}(:[0-9]{1,5})?(\/.*)?$"

	ip_pattern=re.compile(IP_REGEX)
	url_pattern=re.compile(URL_REGEX)
	formattedString = "NA"

	match=ip_pattern.search(s)
	if match:
		ip_group = match.group(1)
		port_group = match.group(2)
		if (match.group(2) == "80"):
			formattedString = "http://" + ip_group
		elif (match.group(2) == "443"):
			formattedString = "https://" + ip_group
		else:
			formattedString = ip_group
		return formattedString

	match=url_pattern.search(s)
	if match:
		formattedString	= match.group(0)
		return formattedString

	return formattedString


#====================

filename=args.file
lines=read_file(filename)
final_list = {'url': 'image'}
dirName = create_dir()


for line in lines:
	print("Checking {}".format(line))

	line = checkFormat(line)

	try:
		driver.get(line)
		imageFileName=str(line) 
		imageFileName = dirName + "/" + re.sub('[^A-Za-z0-9]+', '', imageFileName) + "_screenshot.png"
		time.sleep(2)
		screenshot =  driver.save_screenshot(imageFileName)
		print(".\n")
		final_list[line] = imageFileName
	except Exception as e:
		print("Error {}. Moving on...".format(e))
		final_list[line] = "Timeout"


driver.quit()


document = Document()

document.add_heading('Web site Check', 0)

del final_list['url']

for url, image in final_list.items():
	document.add_paragraph('URL: ' + url + '  :', style='List Bullet')
	if image == "Timeout":
		document.add_paragraph('** timeout **')
	else:
		document.add_picture(image, width=Inches(5.9))

	document.add_page_break()

print(".. writing " + dirName + "/" + 'report.docx') 
document.save(dirName + "/" + 'report.docx')

print("Done")
